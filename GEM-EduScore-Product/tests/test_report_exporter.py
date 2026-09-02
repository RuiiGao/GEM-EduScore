from __future__ import annotations

import unittest
from io import BytesIO

from docx import Document
from pypdf import PdfReader

from modules.comparison import compare_projects
from modules.report_exporter import (
    generate_comparison_docx,
    generate_comparison_pdf,
    generate_docx_report,
    generate_pdf_report,
)
from modules.report_schema import EvaluationPayload, localize_payload, prepare_dashboard
from tests.test_llm_client import make_payload
from tests.test_localization import make_bilingual_payload


class ReportExporterTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        payload = EvaluationPayload.model_validate(make_payload())
        cls.dashboard = prepare_dashboard(payload, "en")
        cls.wiki_pages = [
            {
                "title": "Example Education Wiki",
                "url": "https://2025.igem.wiki/example/education",
            }
        ]

    def test_pdf_export_is_readable_and_paginated(self) -> None:
        result = generate_pdf_report(self.dashboard, "en", self.wiki_pages)

        self.assertTrue(result.startswith(b"%PDF"))
        reader = PdfReader(BytesIO(result))
        self.assertGreaterEqual(len(reader.pages), 2)
        self.assertEqual(reader.metadata.title, self.dashboard["report_title"])

    def test_docx_export_is_editable_and_contains_score_table(self) -> None:
        result = generate_docx_report(self.dashboard, "en", self.wiki_pages)

        self.assertTrue(result.startswith(b"PK"))
        document = Document(BytesIO(result))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn(self.dashboard["report_title"], text)
        self.assertIn("Overall Assessment", text)
        self.assertGreaterEqual(len(document.tables), 2)
        self.assertEqual(document.tables[0].cell(0, 0).text, "Education Design Score")

    def test_chinese_exports_keep_localized_headings(self) -> None:
        localized = localize_payload(make_bilingual_payload(), "zh")
        dashboard = prepare_dashboard(localized, "zh")

        pdf = PdfReader(BytesIO(generate_pdf_report(dashboard, "zh")))
        pdf_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        self.assertIn("总体评估", pdf_text)

        document = Document(BytesIO(generate_docx_report(dashboard, "zh")))
        docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("总体评估", docx_text)

    def test_comparison_exports_include_both_projects_and_chinese_headings(self) -> None:
        localized = localize_payload(make_bilingual_payload(), "zh")
        dashboard_a = prepare_dashboard(localized, "zh")
        dashboard_b = {**dashboard_a, "dimensions": [dict(item) for item in dashboard_a["dimensions"]]}
        dashboard_b["dimensions"][0]["score"] = 6
        comparison = compare_projects(dashboard_a, dashboard_b, "项目甲", "项目乙", "zh")

        pdf_bytes = generate_comparison_pdf(comparison, dashboard_a, dashboard_b, "zh")
        self.assertTrue(pdf_bytes.startswith(b"%PDF"))
        pdf = PdfReader(BytesIO(pdf_bytes))
        pdf_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        self.assertGreaterEqual(len(pdf.pages), 2)
        self.assertIn("十维对照", pdf_text)
        self.assertIn("项目甲", pdf_text)
        self.assertIn("项目乙", pdf_text)

        docx_bytes = generate_comparison_docx(comparison, dashboard_a, dashboard_b, "zh")
        self.assertTrue(docx_bytes.startswith(b"PK"))
        document = Document(BytesIO(docx_bytes))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("十维对照", text)
        self.assertIn("项目甲", text)
        self.assertIn("项目乙", text)
        self.assertGreaterEqual(len(document.tables), 2)


if __name__ == "__main__":
    unittest.main()
