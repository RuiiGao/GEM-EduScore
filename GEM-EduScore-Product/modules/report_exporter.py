"""Polished PDF and editable Word exports for structured GEM-EduScore results."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape, quoteattr

from .localization import OutputLanguage, localized_text


INK = "17233D"
MUTED = "60708C"
INDIGO = "4F46E5"
TEAL = "0F9F8F"
LIGHT = "F3F5FA"
LINE = "DDE3EE"


def generate_pdf_report(
    dashboard: dict[str, Any],
    language: OutputLanguage,
    wiki_pages: list[dict[str, Any]] | None = None,
) -> bytes:
    """Create a paginated, Unicode-safe PDF report entirely in memory."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        KeepTogether,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    t = lambda zh, en: localized_text(zh, en, language)
    font_regular, font_bold = _register_pdf_fonts()
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.78 * inch,
        bottomMargin=0.72 * inch,
        title=str(dashboard.get("report_title", "GEM-EduScore Report")),
        author="GEM-EduScore",
        subject=t("教育实践证据诊断报告", "Education practice evidence diagnostic report"),
    )
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "GEMTitle",
            parent=base["Title"],
            fontName=font_bold,
            fontSize=22,
            leading=28,
            textColor=colors.HexColor(f"#{INK}"),
            alignment=TA_LEFT,
            spaceAfter=10,
        ),
        "kicker": ParagraphStyle(
            "GEMKicker",
            parent=base["Normal"],
            fontName=font_bold,
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor(f"#{INDIGO}"),
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "GEMBody",
            parent=base["BodyText"],
            fontName=font_regular,
            fontSize=10,
            leading=16,
            textColor=colors.HexColor(f"#{INK}"),
            spaceAfter=7,
        ),
        "muted": ParagraphStyle(
            "GEMMuted",
            parent=base["BodyText"],
            fontName=font_regular,
            fontSize=8.5,
            leading=13,
            textColor=colors.HexColor(f"#{MUTED}"),
            spaceAfter=5,
        ),
        "h1": ParagraphStyle(
            "GEMH1",
            parent=base["Heading1"],
            fontName=font_bold,
            fontSize=15,
            leading=20,
            textColor=colors.HexColor(f"#{INK}"),
            spaceBefore=14,
            spaceAfter=8,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "GEMH2",
            parent=base["Heading2"],
            fontName=font_bold,
            fontSize=11.5,
            leading=16,
            textColor=colors.HexColor(f"#{INDIGO}"),
            spaceBefore=10,
            spaceAfter=5,
            keepWithNext=True,
        ),
        "table": ParagraphStyle(
            "GEMTable",
            parent=base["BodyText"],
            fontName=font_regular,
            fontSize=7.8,
            leading=10.5,
            textColor=colors.HexColor(f"#{INK}"),
        ),
        "table_head": ParagraphStyle(
            "GEMTableHead",
            parent=base["BodyText"],
            fontName=font_bold,
            fontSize=7.6,
            leading=9.5,
            textColor=colors.white,
            alignment=TA_CENTER,
        ),
        "quote": ParagraphStyle(
            "GEMQuote",
            parent=base["BodyText"],
            fontName=font_regular,
            fontSize=9,
            leading=14,
            textColor=colors.HexColor(f"#{MUTED}"),
            leftIndent=10,
            borderColor=colors.HexColor(f"#{LINE}"),
            borderWidth=0,
            borderPadding=(4, 8, 4, 8),
            backColor=colors.HexColor("#F8FAFD"),
            spaceAfter=6,
        ),
    }

    def paragraph(value: Any, style: str = "body", *, markup: bool = False):
        content = str(value or "") if markup else _safe_markup(value)
        return Paragraph(content, styles[style])

    def bullet(value: Any):
        return Paragraph(f"&#8226;&nbsp;&nbsp;{_safe_markup(value)}", styles["body"])

    story: list[Any] = [
        Paragraph(escape(t("GEM-EDUSCORE 教育评估报告", "GEM-EDUSCORE EDUCATION EVALUATION")), styles["kicker"]),
        paragraph(dashboard.get("report_title", "GEM-EduScore Report"), "title"),
        paragraph(
            f"<b>{escape(t('评价对象', 'Evaluation Object'))}:</b> {_safe_markup(dashboard.get('practice_name', ''))}<br/>"
            f"<b>{escape(t('团队 / 年份', 'Team / Year'))}:</b> {_safe_markup(dashboard.get('team', ''))} / {_safe_markup(dashboard.get('year', ''))}<br/>"
            f"<b>{escape(t('评价范围', 'Evaluation Scope'))}:</b> {_safe_markup(dashboard.get('evaluation_scope', ''))}",
            "muted",
            markup=True,
        ),
        Spacer(1, 5),
    ]

    summary_data = [
        [
            paragraph(t("教育设计得分", "Education Design Score"), "table_head"),
            paragraph(t("证据覆盖率", "Evidence Coverage"), "table_head"),
            paragraph(t("最强维度", "Strongest Dimension"), "table_head"),
            paragraph(t("优先改进", "Priority Focus"), "table_head"),
        ],
        [
            paragraph(f"{dashboard['design_score']:.1f} / 100", "body"),
            paragraph(f"{dashboard['evidence_coverage']:.0f}%", "body"),
            paragraph(f"{dashboard['strongest_dimension']['id']} · {dashboard['strongest_dimension']['name']}", "body"),
            paragraph(f"{dashboard['priority_dimension']['id']} · {dashboard['priority_dimension']['name']}", "body"),
        ],
    ]
    summary_table = Table(summary_data, colWidths=[1.625 * inch] * 4, repeatRows=1)
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{INDIGO}")),
                ("BACKGROUND", (0, 1), (-1, 1), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{LINE}")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 1), (-1, 1), "CENTER"),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.extend(
        [
            summary_table,
            Spacer(1, 10),
            paragraph(t("本报告为 GEM-EduScore 诊断结果，不是 iGEM 官方评分或排名。", "This is a GEM-EduScore diagnostic report, not an official iGEM score or ranking."), "quote"),
            paragraph(t("总体评估", "Overall Assessment"), "h1"),
            paragraph(dashboard.get("summary", "")),
            paragraph(t("教育实践概览", "Education Practice Overview"), "h1"),
        ]
    )
    for heading, key in (
        (t("目标受众", "Target Audiences"), "audiences"),
        (t("教育目标", "Education Goals"), "goals"),
        (t("主要活动 / 记录", "Main Activities / Records"), "activities"),
    ):
        story.append(paragraph(heading, "h2"))
        story.extend(bullet(item) for item in dashboard.get(key, []))

    story.extend([paragraph(t("证据概况", "Evidence Profile"), "h1"), paragraph(t("有力证据", "Strong Evidence"), "h2")])
    for item in dashboard.get("evidence_profile", {}).get("strong_evidence", []):
        story.append(
            KeepTogether(
                [
                    paragraph(f"<b>{_safe_markup(item.get('record_id', ''))} · {_safe_markup(item.get('strength', ''))}</b> - {_safe_markup(item.get('statement', ''))}", markup=True),
                    paragraph(f"<b>{escape(t('证据原文（原始语言）', 'Source quote'))}:</b><br/>&#8220;{_safe_markup(item.get('source_quote', ''))}&#8221;", "quote", markup=True),
                ]
            )
        )
    story.append(paragraph(t("缺失证据", "Missing Evidence"), "h2"))
    story.extend(bullet(item) for item in dashboard.get("evidence_profile", {}).get("missing_evidence", []))

    story.extend([PageBreak(), paragraph(t("十维量规评价", "Ten-dimension Rubric Evaluation"), "h1")])
    dimension_rows = [[
        paragraph(t("维度", "Dimension"), "table_head"),
        paragraph(t("得分", "Score"), "table_head"),
        paragraph(t("权重", "Weight"), "table_head"),
        paragraph(t("证据", "Evidence"), "table_head"),
        paragraph(t("加权贡献", "Contribution"), "table_head"),
    ]]
    for item in dashboard.get("dimensions", []):
        dimension_rows.append(
            [
                paragraph(f"{item['id']} · {item['name']}", "table"),
                paragraph(f"{item['score']} / 6", "table"),
                paragraph(f"{item['weight']}%", "table"),
                paragraph(item["evidence_strength"], "table"),
                paragraph(f"{item['contribution']:.1f}", "table"),
            ]
        )
    dimension_table = Table(
        dimension_rows,
        colWidths=[2.8 * inch, 0.7 * inch, 0.7 * inch, 0.9 * inch, 1.4 * inch],
        repeatRows=1,
    )
    dimension_table.setStyle(_pdf_table_style())
    story.extend([dimension_table, Spacer(1, 8), paragraph(t("维度诊断", "Dimension-level Diagnosis"), "h1")])
    for item in dashboard.get("dimensions", []):
        quote_text = "<br/>".join(f"&#8220;{_safe_markup(quote)}&#8221;" for quote in item.get("evidence_quotes", []))
        if not quote_text:
            quote_text = escape(t("未发现直接证据", "No direct evidence found"))
        story.extend(
            [
                paragraph(f"{item['id']} · {item['name']} - {item['score']}/6 · {item['evidence_strength']}", "h2"),
                paragraph(f"<b>{escape(t('评价', 'Evaluation'))}:</b> {_safe_markup(item.get('reason', ''))}", markup=True),
                paragraph(f"<b>{escape(t('为何不能更高', 'Why not higher'))}:</b> {_safe_markup(item.get('why_not_higher', ''))}", markup=True),
                paragraph(f"<b>{escape(t('下一步', 'Next move'))}:</b> {_safe_markup(item.get('improvement', ''))}", markup=True),
                paragraph(f"<b>{escape(t('证据原文（保留原始语言）', 'Source evidence excerpts'))}:</b><br/>{quote_text}", "quote", markup=True),
            ]
        )

    story.extend([paragraph(t("基准比较", "Benchmark Comparison"), "h1"), paragraph(f"<b>{escape(t('比较基准', 'Benchmark'))}:</b> {_safe_markup(dashboard.get('benchmark_name', ''))}", markup=True)])
    story.extend(bullet(item) for item in dashboard.get("benchmark_similarities", []))
    for gap in dashboard.get("benchmark_gaps", []):
        story.extend(
            [
                paragraph(f"{gap.get('dimension', '')} · {gap.get('priority', '')}", "h2"),
                paragraph(f"<b>{escape(t('差距', 'Gap'))}:</b> {_safe_markup(gap.get('gap', ''))}", markup=True),
                paragraph(f"<b>{escape(t('改进机会', 'Opportunity'))}:</b> {_safe_markup(gap.get('opportunity', ''))}", markup=True),
            ]
        )

    story.append(paragraph(t("行动路线", "Action Roadmap"), "h1"))
    recommendations = dashboard.get("recommendations", {})
    for heading, key in (
        (t("短期行动", "Short-term Actions"), "short_term"),
        (t("中期策略", "Medium-term Strategies"), "medium_term"),
        (t("长期发展", "Long-term Development"), "long_term"),
    ):
        story.append(paragraph(heading, "h2"))
        story.extend(bullet(item) for item in recommendations.get(key, []))
    story.extend([paragraph(t("结论", "Conclusion"), "h1"), paragraph(dashboard.get("conclusion", ""))])

    if wiki_pages:
        story.append(paragraph(t("Wiki 来源页面", "Wiki Source Pages"), "h1"))
        for page in wiki_pages:
            title = _safe_markup(page.get("title", "Wiki page"))
            url = str(page.get("url", ""))
            if url.startswith(("https://", "http://")):
                story.append(Paragraph(f'&#8226;&nbsp;&nbsp;<link href={quoteattr(url)} color="#{INDIGO}">{title}</link>', styles["body"]))

    def draw_page(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor(f"#{LINE}"))
        canvas.setLineWidth(0.5)
        canvas.line(inch, 0.55 * inch, 7.5 * inch, 0.55 * inch)
        canvas.setFont(font_regular, 8)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawString(inch, 0.36 * inch, "GEM-EduScore")
        canvas.drawRightString(7.5 * inch, 0.36 * inch, f"{t('第', 'Page')} {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    return output.getvalue()


def generate_docx_report(
    dashboard: dict[str, Any],
    language: OutputLanguage,
    wiki_pages: list[dict[str, Any]] | None = None,
) -> bytes:
    """Create an editable Word report with deterministic styles and tables."""
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    t = lambda zh, en: localized_text(zh, en, language)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    _configure_docx_styles(document, language)
    _set_docx_header_footer(section, t("教育实践证据诊断报告", "Education Practice Evidence Diagnostic Report"))
    document.core_properties.title = str(dashboard.get("report_title", "GEM-EduScore Report"))
    document.core_properties.author = "GEM-EduScore"
    document.core_properties.subject = t("教育实践证据诊断报告", "Education practice evidence diagnostic report")

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(5)
    run = kicker.add_run(t("GEM-EDUSCORE 教育评估报告", "GEM-EDUSCORE EDUCATION EVALUATION"))
    _format_docx_run(run, 9, INDIGO, bold=True)
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(9)
    title.paragraph_format.keep_with_next = True
    _format_docx_run(title.add_run(str(dashboard.get("report_title", "GEM-EduScore Report"))), 24, INK, bold=True)
    _add_labeled_docx_paragraph(document, t("评价对象", "Evaluation Object"), dashboard.get("practice_name", ""), muted=True)
    _add_labeled_docx_paragraph(document, t("团队 / 年份", "Team / Year"), f"{dashboard.get('team', '')} / {dashboard.get('year', '')}", muted=True)
    _add_labeled_docx_paragraph(document, t("评价范围", "Evaluation Scope"), dashboard.get("evaluation_scope", ""), muted=True)

    score_table = document.add_table(rows=2, cols=4)
    score_table.autofit = False
    score_labels = [
        t("教育设计得分", "Education Design Score"),
        t("证据覆盖率", "Evidence Coverage"),
        t("最强维度", "Strongest Dimension"),
        t("优先改进", "Priority Focus"),
    ]
    score_values = [
        f"{dashboard['design_score']:.1f} / 100",
        f"{dashboard['evidence_coverage']:.0f}%",
        f"{dashboard['strongest_dimension']['id']} · {dashboard['strongest_dimension']['name']}",
        f"{dashboard['priority_dimension']['id']} · {dashboard['priority_dimension']['name']}",
    ]
    for column, (label, value) in enumerate(zip(score_labels, score_values)):
        _set_docx_cell_text(score_table.cell(0, column), label, bold=True, color="FFFFFF", align="center")
        _set_docx_cell_text(score_table.cell(1, column), value, align="center")
        _shade_docx_cell(score_table.cell(0, column), INDIGO)
    _set_docx_table_geometry(score_table, [2340, 2340, 2340, 2340], indent=120)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(9)
    note.paragraph_format.space_after = Pt(8)
    _format_docx_run(note.add_run(t("本报告为 GEM-EduScore 诊断结果，不是 iGEM 官方评分或排名。", "This is a GEM-EduScore diagnostic report, not an official iGEM score or ranking.")), 9.5, MUTED, italic=True)

    document.add_heading(t("总体评估", "Overall Assessment"), level=1)
    document.add_paragraph(str(dashboard.get("summary", "")))
    document.add_heading(t("教育实践概览", "Education Practice Overview"), level=1)
    for heading, key in (
        (t("目标受众", "Target Audiences"), "audiences"),
        (t("教育目标", "Education Goals"), "goals"),
        (t("主要活动 / 记录", "Main Activities / Records"), "activities"),
    ):
        document.add_heading(heading, level=2)
        _add_docx_bullets(document, dashboard.get(key, []))

    document.add_heading(t("证据概况", "Evidence Profile"), level=1)
    document.add_heading(t("有力证据", "Strong Evidence"), level=2)
    for item in dashboard.get("evidence_profile", {}).get("strong_evidence", []):
        _add_labeled_docx_paragraph(document, f"{item.get('record_id', '')} · {item.get('strength', '')}", item.get("statement", ""))
        quote = document.add_paragraph()
        quote.paragraph_format.left_indent = Inches(0.18)
        quote.paragraph_format.space_after = Pt(6)
        _format_docx_run(quote.add_run(f"{t('证据原文（原始语言）', 'Source quote')}: “{item.get('source_quote', '')}”"), 9.5, MUTED, italic=True)
    document.add_heading(t("缺失证据", "Missing Evidence"), level=2)
    _add_docx_bullets(document, dashboard.get("evidence_profile", {}).get("missing_evidence", []))

    document.add_page_break()
    document.add_heading(t("十维量规评价", "Ten-dimension Rubric Evaluation"), level=1)
    dimension_table = document.add_table(rows=1, cols=5)
    dimension_table.autofit = False
    headers = [t("维度", "Dimension"), t("得分", "Score"), t("权重", "Weight"), t("证据", "Evidence"), t("加权贡献", "Contribution")]
    for index, label in enumerate(headers):
        _set_docx_cell_text(dimension_table.cell(0, index), label, bold=True, color="FFFFFF", align="center")
        _shade_docx_cell(dimension_table.cell(0, index), INDIGO)
    _repeat_docx_table_header(dimension_table.rows[0])
    for item in dashboard.get("dimensions", []):
        cells = dimension_table.add_row().cells
        values = [
            f"{item['id']} · {item['name']}",
            f"{item['score']} / 6",
            f"{item['weight']}%",
            item["evidence_strength"],
            f"{item['contribution']:.1f}",
        ]
        for index, value in enumerate(values):
            _set_docx_cell_text(cells[index], value, align="left" if index == 0 else "center")
    _set_docx_table_geometry(dimension_table, [4032, 1008, 1008, 1296, 2016], indent=120)

    document.add_heading(t("维度诊断", "Dimension-level Diagnosis"), level=1)
    for item in dashboard.get("dimensions", []):
        document.add_heading(f"{item['id']} · {item['name']} - {item['score']}/6 · {item['evidence_strength']}", level=2)
        _add_labeled_docx_paragraph(document, t("评价", "Evaluation"), item.get("reason", ""))
        _add_labeled_docx_paragraph(document, t("为何不能更高", "Why not higher"), item.get("why_not_higher", ""))
        _add_labeled_docx_paragraph(document, t("下一步", "Next move"), item.get("improvement", ""))
        for quote_text in item.get("evidence_quotes", []):
            quote = document.add_paragraph()
            quote.paragraph_format.left_indent = Inches(0.18)
            _format_docx_run(quote.add_run(f"“{quote_text}”"), 9.5, MUTED, italic=True)

    document.add_heading(t("基准比较", "Benchmark Comparison"), level=1)
    _add_labeled_docx_paragraph(document, t("比较基准", "Benchmark"), dashboard.get("benchmark_name", ""))
    _add_docx_bullets(document, dashboard.get("benchmark_similarities", []))
    for gap in dashboard.get("benchmark_gaps", []):
        document.add_heading(f"{gap.get('dimension', '')} · {gap.get('priority', '')}", level=2)
        _add_labeled_docx_paragraph(document, t("差距", "Gap"), gap.get("gap", ""))
        _add_labeled_docx_paragraph(document, t("改进机会", "Opportunity"), gap.get("opportunity", ""))

    document.add_heading(t("行动路线", "Action Roadmap"), level=1)
    recommendations = dashboard.get("recommendations", {})
    for heading, key in (
        (t("短期行动", "Short-term Actions"), "short_term"),
        (t("中期策略", "Medium-term Strategies"), "medium_term"),
        (t("长期发展", "Long-term Development"), "long_term"),
    ):
        document.add_heading(heading, level=2)
        _add_docx_bullets(document, recommendations.get(key, []))
    document.add_heading(t("结论", "Conclusion"), level=1)
    document.add_paragraph(str(dashboard.get("conclusion", "")))

    if wiki_pages:
        document.add_heading(t("Wiki 来源页面", "Wiki Source Pages"), level=1)
        for page in wiki_pages:
            document.add_paragraph(f"{page.get('title', 'Wiki page')} - {page.get('url', '')}", style="List Bullet")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def generate_comparison_pdf(
    comparison: dict[str, Any],
    dashboard_a: dict[str, Any],
    dashboard_b: dict[str, Any],
    language: OutputLanguage,
    sources_a: list[dict[str, Any]] | None = None,
    sources_b: list[dict[str, Any]] | None = None,
) -> bytes:
    """Create a polished, self-contained PDF for a two-project comparison."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    t = lambda zh, en: localized_text(zh, en, language)
    font_regular, font_bold = _register_pdf_fonts()
    label_a = str(comparison["label_a"])
    label_b = str(comparison["label_b"])
    title_text = t(
        f"GEM-EduScore 双项目对比报告：{label_a} vs {label_b}",
        f"GEM-EduScore Project Comparison: {label_a} vs {label_b}",
    )
    visible_title = t("GEM-EduScore 双项目对比报告", "GEM-EduScore Project Comparison Report")
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.78 * inch,
        bottomMargin=0.72 * inch,
        title=title_text,
        author="GEM-EduScore",
        subject=t("双项目教育实践证据对比", "Two-project education evidence comparison"),
    )
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle("ComparisonTitle", parent=base["Title"], fontName=font_bold, fontSize=21, leading=27, textColor=colors.HexColor(f"#{INK}"), alignment=TA_LEFT, spaceAfter=9),
        "subtitle": ParagraphStyle("ComparisonSubtitle", parent=base["Heading2"], fontName=font_bold, fontSize=14, leading=19, textColor=colors.HexColor(f"#{INDIGO}"), alignment=TA_LEFT, spaceAfter=9),
        "kicker": ParagraphStyle("ComparisonKicker", parent=base["Normal"], fontName=font_bold, fontSize=8.5, leading=11, textColor=colors.HexColor(f"#{INDIGO}"), spaceAfter=6),
        "body": ParagraphStyle("ComparisonBody", parent=base["BodyText"], fontName=font_regular, fontSize=9.7, leading=15, textColor=colors.HexColor(f"#{INK}"), spaceAfter=6),
        "muted": ParagraphStyle("ComparisonMuted", parent=base["BodyText"], fontName=font_regular, fontSize=8.3, leading=12, textColor=colors.HexColor(f"#{MUTED}"), spaceAfter=5),
        "h1": ParagraphStyle("ComparisonH1", parent=base["Heading1"], fontName=font_bold, fontSize=15, leading=20, textColor=colors.HexColor(f"#{INK}"), spaceBefore=14, spaceAfter=8, keepWithNext=True),
        "h2": ParagraphStyle("ComparisonH2", parent=base["Heading2"], fontName=font_bold, fontSize=11.2, leading=15, textColor=colors.HexColor(f"#{INDIGO}"), spaceBefore=9, spaceAfter=4, keepWithNext=True),
        "table": ParagraphStyle("ComparisonTable", parent=base["BodyText"], fontName=font_regular, fontSize=7.8, leading=10, textColor=colors.HexColor(f"#{INK}")),
        "table_head": ParagraphStyle("ComparisonTableHead", parent=base["BodyText"], fontName=font_bold, fontSize=7.5, leading=9.2, textColor=colors.white, alignment=TA_CENTER),
        "metric": ParagraphStyle("ComparisonMetric", parent=base["BodyText"], fontName=font_bold, fontSize=13, leading=16, textColor=colors.HexColor(f"#{INK}"), alignment=TA_CENTER),
    }

    def paragraph(value: Any, style: str = "body", *, markup: bool = False):
        content = str(value or "") if markup else _safe_markup(value)
        return Paragraph(content, styles[style])

    story: list[Any] = [
        Paragraph(escape(t("GEM-EDUSCORE 双项目证据对比", "GEM-EDUSCORE PAIRWISE EVIDENCE COMPARISON")), styles["kicker"]),
        paragraph(visible_title, "title"),
        paragraph(f"{label_a} vs {label_b}", "subtitle"),
        paragraph(t("同一量规、同一案例组合、两次独立评价。差异只反映已提供材料中的证据。", "Same rubric, same benchmark portfolio and two independent evaluations. Differences reflect only the supplied evidence."), "muted"),
    ]
    metrics = [
        [paragraph(label_a, "table_head"), paragraph(label_b, "table_head"), paragraph(t("设计得分差 A-B", "Design delta A-B"), "table_head"), paragraph(t("证据覆盖差 A-B", "Evidence delta A-B"), "table_head")],
        [paragraph(f"{dashboard_a['design_score']:.1f} / 100", "metric"), paragraph(f"{dashboard_b['design_score']:.1f} / 100", "metric"), paragraph(f"{comparison['design_delta']:+.1f}", "metric"), paragraph(f"{comparison['evidence_delta']:+.1f} pp", "metric")],
    ]
    metric_table = Table(metrics, colWidths=[1.625 * inch] * 4, repeatRows=1)
    metric_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{INDIGO}")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(f"#{LINE}")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 1), (-1, 1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.extend([
        Spacer(1, 5), metric_table, Spacer(1, 9),
        paragraph(t("本报告为 GEM-EduScore 诊断性对比，不构成 iGEM 官方评分或排名。", "This is a GEM-EduScore diagnostic comparison, not an official iGEM score or ranking."), "muted"),
        paragraph(t("项目概览", "Project Overview"), "h1"),
        paragraph(f"<b>{escape(label_a)}</b><br/>{_safe_markup(dashboard_a.get('summary', ''))}", "body", markup=True),
        paragraph(f"<b>{escape(label_b)}</b><br/>{_safe_markup(dashboard_b.get('summary', ''))}", "body", markup=True),
        PageBreak(),
        paragraph(t("十维对照", "Ten-dimension Comparison"), "h1"),
    ])
    dimension_rows = [[
        paragraph(t("维度", "Dimension"), "table_head"),
        paragraph(label_a, "table_head"),
        paragraph(label_b, "table_head"),
        paragraph("Δ A-B", "table_head"),
        paragraph(t("领先", "Lead"), "table_head"),
    ]]
    for row in comparison["rows"]:
        dimension_rows.append([
            paragraph(f"{row['id']} · {row['name_a']}", "table"),
            paragraph(f"{row['score_a']}/6 · {row['evidence_a']}", "table"),
            paragraph(f"{row['score_b']}/6 · {row['evidence_b']}", "table"),
            paragraph(f"{row['delta']:+d}", "table"),
            paragraph(row["lead"], "table"),
        ])
    dimension_table = Table(dimension_rows, colWidths=[2.65 * inch, 1.0 * inch, 1.0 * inch, 0.65 * inch, 1.2 * inch], repeatRows=1)
    dimension_table.setStyle(_pdf_table_style())
    story.extend([dimension_table, paragraph(t("互相借鉴建议", "Cross-learning Recommendations"), "h1")])
    for index, item in enumerate(comparison.get("recommendations", []), 1):
        story.append(paragraph(f"<b>{index:02d}</b>&nbsp;&nbsp;{_safe_markup(item)}", "body", markup=True))

    story.append(paragraph(t("逐维下一步", "Dimension-level Next Moves"), "h1"))
    for row in comparison["rows"]:
        story.append(KeepTogether([
            paragraph(f"{row['id']} · {row['name_a']}", "h2"),
            paragraph(f"<b>{escape(label_a)} · {escape(t('下一步', 'Next move'))}:</b> {_safe_markup(row['improvement_a'])}", "body", markup=True),
            paragraph(f"<b>{escape(label_b)} · {escape(t('下一步', 'Next move'))}:</b> {_safe_markup(row['improvement_b'])}", "body", markup=True),
        ]))

    source_groups = [(label_a, sources_a or []), (label_b, sources_b or [])]
    if any(pages for _, pages in source_groups):
        story.append(paragraph(t("Wiki 来源页面", "Wiki Source Pages"), "h1"))
        for label, pages in source_groups:
            if not pages:
                continue
            story.append(paragraph(label, "h2"))
            for page in pages:
                title = page.get("title", "Wiki page")
                url = page.get("url", "")
                story.append(paragraph(f"{title} - {url}", "muted"))

    def draw_page(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor(f"#{LINE}"))
        canvas.setLineWidth(0.5)
        canvas.line(inch, 0.55 * inch, 7.5 * inch, 0.55 * inch)
        canvas.setFont(font_regular, 8)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawString(inch, 0.36 * inch, "GEM-EduScore")
        canvas.drawRightString(7.5 * inch, 0.36 * inch, f"{t('第', 'Page')} {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    return output.getvalue()


def generate_comparison_docx(
    comparison: dict[str, Any],
    dashboard_a: dict[str, Any],
    dashboard_b: dict[str, Any],
    language: OutputLanguage,
    sources_a: list[dict[str, Any]] | None = None,
    sources_b: list[dict[str, Any]] | None = None,
) -> bytes:
    """Create an editable Word comparison using the business-brief preset."""
    from docx import Document
    from docx.shared import Inches, Pt

    t = lambda zh, en: localized_text(zh, en, language)
    label_a = str(comparison["label_a"])
    label_b = str(comparison["label_b"])
    title_text = t(
        f"GEM-EduScore 双项目对比报告：{label_a} vs {label_b}",
        f"GEM-EduScore Project Comparison: {label_a} vs {label_b}",
    )
    visible_title = t("GEM-EduScore 双项目对比报告", "GEM-EduScore Project Comparison Report")
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    _configure_docx_styles(document, language)
    _set_docx_header_footer(section, t("双项目教育实践证据对比", "Two-project Education Evidence Comparison"))
    document.core_properties.title = title_text
    document.core_properties.author = "GEM-EduScore"
    document.core_properties.subject = t("双项目教育实践证据对比", "Two-project education evidence comparison")

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(5)
    _format_docx_run(kicker.add_run(t("GEM-EDUSCORE 双项目证据对比", "GEM-EDUSCORE PAIRWISE EVIDENCE COMPARISON")), 9, INDIGO, bold=True)
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    _format_docx_run(title.add_run(visible_title), 23, INK, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle.paragraph_format.keep_with_next = True
    _format_docx_run(subtitle.add_run(f"{label_a} vs {label_b}"), 15, INDIGO, bold=True)
    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(9)
    _format_docx_run(intro.add_run(t("同一量规、同一案例组合、两次独立评价。差异只反映已提供材料中的证据。", "Same rubric, same benchmark portfolio and two independent evaluations. Differences reflect only the supplied evidence.")), 9.5, MUTED)

    metrics = document.add_table(rows=1, cols=3)
    headers = [t("指标", "Metric"), label_a, label_b]
    for index, value in enumerate(headers):
        _set_docx_cell_text(metrics.cell(0, index), value, bold=True, color="FFFFFF", align="center")
        _shade_docx_cell(metrics.cell(0, index), INDIGO)
    for metric, value_a, value_b in (
        (t("教育设计得分", "Education Design Score"), f"{dashboard_a['design_score']:.1f} / 100", f"{dashboard_b['design_score']:.1f} / 100"),
        (t("证据覆盖率", "Evidence Coverage"), f"{dashboard_a['evidence_coverage']:.1f}%", f"{dashboard_b['evidence_coverage']:.1f}%"),
    ):
        cells = metrics.add_row().cells
        for index, value in enumerate((metric, value_a, value_b)):
            _set_docx_cell_text(cells[index], value, align="left" if index == 0 else "center")
    _set_docx_table_geometry(metrics, [3600, 2880, 2880], indent=120)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(8)
    _format_docx_run(note.add_run(t("本报告为 GEM-EduScore 诊断性对比，不构成 iGEM 官方评分或排名。", "This is a GEM-EduScore diagnostic comparison, not an official iGEM score or ranking.")), 9.5, MUTED, italic=True)

    document.add_heading(t("项目概览", "Project Overview"), level=1)
    document.add_heading(label_a, level=2)
    document.add_paragraph(str(dashboard_a.get("summary", "")))
    document.add_heading(label_b, level=2)
    document.add_paragraph(str(dashboard_b.get("summary", "")))

    document.add_page_break()
    document.add_heading(t("十维对照", "Ten-dimension Comparison"), level=1)
    dimensions = document.add_table(rows=1, cols=5)
    dimension_headers = [t("维度", "Dimension"), label_a, label_b, "Δ A-B", t("领先", "Lead")]
    for index, value in enumerate(dimension_headers):
        _set_docx_cell_text(dimensions.cell(0, index), value, bold=True, color="FFFFFF", align="center")
        _shade_docx_cell(dimensions.cell(0, index), INDIGO)
    _repeat_docx_table_header(dimensions.rows[0])
    for row in comparison["rows"]:
        cells = dimensions.add_row().cells
        values = [f"{row['id']} · {row['name_a']}", f"{row['score_a']}/6 · {row['evidence_a']}", f"{row['score_b']}/6 · {row['evidence_b']}", f"{row['delta']:+d}", row["lead"]]
        for index, value in enumerate(values):
            _set_docx_cell_text(cells[index], value, align="left" if index == 0 else "center")
    _set_docx_table_geometry(dimensions, [4032, 1440, 1440, 1008, 1440], indent=120)

    document.add_heading(t("互相借鉴建议", "Cross-learning Recommendations"), level=1)
    _add_docx_bullets(document, comparison.get("recommendations", []))
    document.add_heading(t("逐维下一步", "Dimension-level Next Moves"), level=1)
    for row in comparison["rows"]:
        heading = document.add_heading(f"{row['id']} · {row['name_a']}", level=2)
        heading.paragraph_format.keep_with_next = True
        moves = document.add_paragraph()
        moves.paragraph_format.space_after = Pt(5)
        moves.paragraph_format.keep_together = True
        _format_docx_run(moves.add_run(f"{label_a} · {t('下一步', 'Next move')}: "), 11, INK, bold=True)
        _format_docx_run(moves.add_run(str(row["improvement_a"])), 11, INK)
        moves.add_run().add_break()
        _format_docx_run(moves.add_run(f"{label_b} · {t('下一步', 'Next move')}: "), 11, INK, bold=True)
        _format_docx_run(moves.add_run(str(row["improvement_b"])), 11, INK)

    source_groups = [(label_a, sources_a or []), (label_b, sources_b or [])]
    if any(pages for _, pages in source_groups):
        document.add_heading(t("Wiki 来源页面", "Wiki Source Pages"), level=1)
        for label, pages in source_groups:
            if not pages:
                continue
            document.add_heading(label, level=2)
            for page in pages:
                document.add_paragraph(f"{page.get('title', 'Wiki page')} - {page.get('url', '')}", style="List Bullet")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _safe_markup(value: Any) -> str:
    return escape(str(value or "")).replace("\n", "<br/>")


def _register_pdf_fonts() -> tuple[str, str]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        (r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\msyhbd.ttc"),
        (r"C:\Windows\Fonts\simsun.ttc", r"C:\Windows\Fonts\simsun.ttc"),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"),
        ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
    ]
    for regular_path, bold_path in candidates:
        if not Path(regular_path).exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont("GEMCJK", regular_path, subfontIndex=0))
            if Path(bold_path).exists():
                pdfmetrics.registerFont(TTFont("GEMCJKBold", bold_path, subfontIndex=0))
            else:
                pdfmetrics.registerFont(TTFont("GEMCJKBold", regular_path, subfontIndex=0))
            return "GEMCJK", "GEMCJKBold"
        except Exception:
            continue
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    except Exception:
        pass
    return "STSong-Light", "STSong-Light"


def _pdf_table_style():
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle

    return TableStyle(
        [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{INDIGO}")),
            ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor(f"#{LINE}")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFD")]),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ]
    )


def _configure_docx_styles(document, language: OutputLanguage) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor, Inches

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in (
        ("Heading 1", 16, INDIGO, 16, 8),
        ("Heading 2", 13, INDIGO, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    bullet = document.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet.font.size = Pt(11)
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167


def _format_docx_run(run, size: float, color: str, *, bold: bool = False, italic: bool = False) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _add_labeled_docx_paragraph(document, label: str, value: Any, *, muted: bool = False):
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3 if muted else 6)
    _format_docx_run(paragraph.add_run(f"{label}: "), 9.5 if muted else 11, MUTED if muted else INK, bold=True)
    _format_docx_run(paragraph.add_run(str(value or "")), 9.5 if muted else 11, MUTED if muted else INK)
    return paragraph


def _add_docx_bullets(document, values: list[Any]) -> None:
    for value in values:
        document.add_paragraph(str(value), style="List Bullet")


def _set_docx_header_footer(section, header_text: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _format_docx_run(header.add_run(f"GEM-EduScore  |  {header_text}"), 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _format_docx_run(footer.add_run("GEM-EduScore  |  "), 8.5, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def _set_docx_cell_text(cell, value: Any, *, bold: bool = False, color: str = INK, align: str = "left") -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    _format_docx_run(paragraph.add_run(str(value or "")), 9, color, bold=bold)


def _shade_docx_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell._tc.get_or_add_tcPr().append(shading)
    shading.set(qn("w:fill"), fill)


def _set_docx_table_geometry(table, widths_dxa: list[int], *, indent: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    total = sum(widths_dxa)
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(total))
    indentation = properties.find(qn("w:tblInd"))
    if indentation is None:
        indentation = OxmlElement("w:tblInd")
        properties.append(indentation)
    indentation.set(qn("w:type"), "dxa")
    indentation.set(qn("w:w"), str(indent))
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(widths_dxa[index]))
            margins = cell._tc.get_or_add_tcPr().find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell._tc.get_or_add_tcPr().append(margins)
            for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def _repeat_docx_table_header(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)
